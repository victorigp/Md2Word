#!/usr/bin/env python
# -*- coding: utf-8 -*-
import sys, os, re, json, argparse, subprocess, tempfile, shutil
from PIL import Image as PilImage
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

S_H1, S_H2, S_H3 = "H1", "H2", "H3"
S_BODY, S_BULLET, S_CODE, S_CAPTION, S_FOOTER = "Body", "Bullet", "Code", "Caption", "Footer"
STYLE_DEFAULTS = {
    S_H1: "Heading 1", S_H2: "Heading 2", S_H3: "Heading 3",
    S_BODY: "Normal", S_BULLET: "List Bullet", S_CODE: "Code",
    S_CAPTION: "Caption", S_FOOTER: "Footer",
}
SETTINGS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "Settings.json")
MAX_IMG_WIDTH_INCHES  = 6.0
MAX_IMG_HEIGHT_INCHES = 7.5


# ---------------------------------------------------------------------------
# Settings helpers
# ---------------------------------------------------------------------------

def load_settings():
    if os.path.exists(SETTINGS_FILE):
        with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    return {}

def get_style(settings, role):
    return settings.get("Styles", {}).get(role, STYLE_DEFAULTS.get(role, "Normal"))

def style_exists(doc, style_name):
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        return False

def safe_style(doc, settings, role):
    name = get_style(settings, role)
    if style_exists(doc, name):
        return name
    fallback = STYLE_DEFAULTS.get(role, "Normal")
    return fallback if style_exists(doc, fallback) else "Normal"


# ---------------------------------------------------------------------------
# Markdown parser
# ---------------------------------------------------------------------------

def parse_markdown(md_path):
    with open(md_path, "r", encoding="utf-8-sig") as f:
        lines = f.readlines()
    elements = []
    i = 0
    base_dir = os.path.dirname(os.path.abspath(md_path))
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip("\n")
        # Code block
        if stripped.startswith("```"):
            lang = stripped[3:].strip().lower()
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip("\n").startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            i += 1
            code = "\n".join(code_lines)
            if lang == "mermaid":
                elements.append({"type": "mermaid", "codigo": code})
            else:
                elements.append({"type": "code_block", "codigo": code, "lenguaje": lang})
            continue
        # Headings
        m = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if m:
            level = len(m.group(1))
            heading_text = re.sub(r'^\d+(?:\.\d+)*\.?\s+', '', m.group(2).strip())
            i += 1
            normalized = heading_text.lower()
            is_toc = any(kw in normalized for kw in
                         ["ndice", "index", "contenido", "table of content"])
            if is_toc:
                elements.append({"type": "toc_heading", "texto": heading_text})
                elements.append({"type": "toc"})
                while i < len(lines):
                    peek = lines[i].rstrip("\n")
                    ps = peek.strip()
                    if not ps:
                        i += 1; continue
                    if re.match(r'^\s*---+\s*$', ps): break
                    if re.match(r'^#{1,6}\s+', ps): break
                    if re.match(r'^\d+\.\s+', ps): i += 1; continue
                    if re.match(r'^[-*+]\s+', ps): i += 1; continue
                    if peek.startswith("   ") or peek.startswith("\t"):
                        i += 1; continue
                    break
            else:
                elements.append({"type": "h{}".format(level), "texto": heading_text})
            continue
        # Separator
        if re.match(r'^\s*---+\s*$', stripped) or re.match(r'^\s*\*\*\*+\s*$', stripped):
            elements.append({"type": "separator"})
            i += 1; continue
        # Image
        m = re.match(r'^!\[([^\]]*)\]\(([^)"\s]+)(?:\s+"[^"]*")?\)', stripped)
        if m:
            alt = m.group(1)
            ruta = m.group(2)
            if not os.path.isabs(ruta):
                ruta = os.path.normpath(os.path.join(base_dir, ruta))
            elements.append({"type": "image", "ruta": ruta, "alt": alt})
            i += 1; continue
        # Table
        if "|" in stripped and i + 1 < len(lines):
            ns_line = lines[i + 1].rstrip("\n")
            if re.match(r'^\|?\s*[-:]+[-| :]*\s*\|?\s*$', ns_line):
                headers = [c.strip() for c in stripped.strip().strip("|").split("|")]
                rows = []
                i += 2
                while i < len(lines) and "|" in lines[i] and lines[i].strip():
                    rows.append([c.strip() for c in
                                 lines[i].strip().strip("|").split("|")])
                    i += 1
                elements.append({"type": "table", "headers": headers, "rows": rows})
                continue
        # Numbered list
        m = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if m:
            elements.append({"type": "numbered", "texto": m.group(2).strip(),
                              "numero": int(m.group(1))})
            i += 1; continue
        # Bullet
        m = re.match(r'^(\s*)([-*+])\s+(.+)$', stripped)
        if m:
            nivel = 1 if len(m.group(1)) >= 2 else 0
            elements.append({"type": "bullet", "texto": m.group(3).strip(),
                              "nivel": nivel})
            i += 1; continue
        # Text
        if stripped.strip():
            elements.append({"type": "text", "texto": stripped.strip()})
        i += 1
    return elements


# ---------------------------------------------------------------------------
# Inline formatting
# ---------------------------------------------------------------------------

INLINE_PATTERN = re.compile(
    r'(\*\*(.+?)\*\*|__(.+?)__|'
    r'\*(.+?)\*|_(.+?)_|'
    r'~~(.+?)~~|`(.+?)`|'
    r'\[([^\]]+)\]\([^)]+\))'
)

def aplicar_inline(paragraph, texto):
    last_end = 0
    for m in INLINE_PATTERN.finditer(texto):
        if m.start() > last_end:
            paragraph.add_run(texto[last_end:m.start()])
        if m.group(2) is not None:
            paragraph.add_run(m.group(2)).bold = True
        elif m.group(3) is not None:
            paragraph.add_run(m.group(3)).bold = True
        elif m.group(4) is not None:
            paragraph.add_run(m.group(4)).italic = True
        elif m.group(5) is not None:
            paragraph.add_run(m.group(5)).italic = True
        elif m.group(6) is not None:
            paragraph.add_run(m.group(6)).font.strike = True
        elif m.group(7) is not None:
            run = paragraph.add_run(m.group(7))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            try:
                shd = parse_xml(
                    '<w:shd {} w:fill="E0E0E0" w:val="clear"/>'.format(nsdecls('w')))
                run._element.get_or_add_rPr().append(shd)
            except Exception:
                pass
        elif m.group(8) is not None:
            paragraph.add_run(m.group(8))
        last_end = m.end()
    if last_end < len(texto):
        paragraph.add_run(texto[last_end:])
    if last_end == 0 and not paragraph.runs:
        paragraph.add_run(texto)


# ---------------------------------------------------------------------------
# Image scaling
# ---------------------------------------------------------------------------

def get_scaled_width(ruta):
    try:
        with PilImage.open(ruta) as img:
            w_px, h_px = img.size
            if w_px == 0 or h_px == 0:
                return MAX_IMG_WIDTH_INCHES
            dpi_info = img.info.get("dpi", (96, 96))
            if isinstance(dpi_info, (int, float)):
                dpi_info = (dpi_info, dpi_info)
            dpi_x = float(dpi_info[0]) if float(dpi_info[0]) > 0 else 96.0
            dpi_y = float(dpi_info[1]) if float(dpi_info[1]) > 0 else 96.0
            w_nat = w_px / dpi_x
            h_nat = h_px / dpi_y
            return min(min(w_nat, MAX_IMG_WIDTH_INCHES),
                       MAX_IMG_HEIGHT_INCHES * (w_nat / h_nat))
    except Exception:
        return MAX_IMG_WIDTH_INCHES


# ---------------------------------------------------------------------------
# Mermaid rendering
# ---------------------------------------------------------------------------

def find_mmdc():
    if sys.platform == "win32":
        for cmd in ["mmdc.cmd", "mmdc"]:
            p = shutil.which(cmd)
            if p:
                return p
    return "mmdc"

def render_mermaid(code):
    tmp_mmd = None
    try:
        tmp_mmd = tempfile.NamedTemporaryFile(
            suffix=".mmd", delete=False, mode="w", encoding="utf-8")
        tmp_mmd.write(code)
        tmp_mmd.close()
        tmp_png = tmp_mmd.name.replace(".mmd", ".png")
        result = subprocess.run(
            [find_mmdc(), "-i", tmp_mmd.name, "-o", tmp_png,
             "-b", "white", "--scale", "2"],
            capture_output=True, text=True, timeout=30,
            shell=(sys.platform == "win32"))
        if result.returncode == 0 and os.path.exists(tmp_png):
            return tmp_png
        print("[AVISO] Mermaid failed: {}".format(result.stderr.strip()))
        return None
    except FileNotFoundError:
        print("[AVISO] mmdc not found."); return None
    except subprocess.TimeoutExpired:
        print("[AVISO] Mermaid timeout."); return None
    except Exception as e:
        print("[AVISO] Mermaid error: {}".format(e)); return None
    finally:
        if tmp_mmd and os.path.exists(tmp_mmd.name):
            try:
                os.unlink(tmp_mmd.name)
            except OSError:
                pass


# ---------------------------------------------------------------------------
# Bookmark helpers
# ---------------------------------------------------------------------------

def _make_bookmark_name(text, idx):
    safe = re.sub(r'[^a-zA-Z0-9]', '_', text)
    safe = re.sub(r'_+', '_', safe).strip('_')
    if not safe or not safe[0].isalpha():
        safe = 'H' + safe
    return '_Toc{:04d}_{}'.format(idx, safe[:30])

def _assign_bookmarks(elements):
    bm_id = 100
    for elem in elements:
        if elem['type'] in ('h2', 'h3'):
            elem['_bookmark'] = _make_bookmark_name(elem['texto'], bm_id)
            elem['_bm_id'] = bm_id
            bm_id += 1

def _add_bookmark_to_para(para_elem, bm_name, bm_id):
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:id'), str(bm_id))
    bm_start.set(qn('w:name'), bm_name)
    bm_end = OxmlElement('w:bookmarkEnd')
    bm_end.set(qn('w:id'), str(bm_id))
    para_elem.insert(0, bm_start)
    para_elem.append(bm_end)


# ---------------------------------------------------------------------------
# TOC helpers
# ---------------------------------------------------------------------------

def _toc_entry_style(doc, level):
    candidates = {
        1: ["toc 1", "TOC 1", "TDC 1", "TDC1"],
        2: ["toc 2", "TOC 2", "TDC 2", "TDC2"],
    }
    for name in candidates.get(level, []):
        try:
            return doc.styles[name].style_id
        except KeyError:
            pass
    return "Normal"


def _h1_style_xml(doc, settings):
    """Return (rpr_xml, before_twips, after_twips) by reading the H1 style
    from the template, resolving inheritance, so the TOC title matches exactly.
    """
    from lxml import etree

    h1_name = get_style(settings, S_H1)
    style = None
    for name in [h1_name, "Heading 1"]:
        try:
            style = doc.styles[name]
            break
        except KeyError:
            pass
    if style is None:
        return "", "240", "120"

    # Walk base style chain, child properties override parent
    rpr_parts = {}
    before = after = None

    def collect(s):
        if s is None:
            return
        if s.base_style:
            collect(s.base_style)
        elem = s.element
        # run properties
        rpr = elem.find(qn('w:rPr'))
        if rpr is not None:
            for child in rpr:
                rpr_parts[child.tag] = etree.tostring(child, encoding='unicode')
        # paragraph spacing
        pPr = elem.find(qn('w:pPr'))
        if pPr is not None:
            spc = pPr.find(qn('w:spacing'))
            if spc is not None:
                nonlocal before, after
                b = spc.get(qn('w:before'))
                a = spc.get(qn('w:after'))
                if b is not None:
                    before = b
                if a is not None:
                    after = a

    collect(style)

    rpr_xml = ""
    if rpr_parts:
        rpr_xml = (
            '<w:rPr xmlns:w="http://schemas.openxmlformats.org/'
            'wordprocessingml/2006/main">'
            + "".join(rpr_parts.values())
            + "</w:rPr>"
        )

    return rpr_xml, (before or "240"), (after or "120")


def _update_sdt_toc(doc, elements, settings=None):
    if settings is None:
        settings = {}

    body = doc.element.body
    toc_sdt = None
    for child in list(body):
        if child.tag != qn('w:sdt'):
            continue
        # 1) Preferred: SDT doc-part gallery marks TOC blocks
        galleries = [g.get(qn('w:val')) or '' for g in child.iter(qn('w:docPartGallery'))]
        if any('table of contents' in g.lower() for g in galleries):
            toc_sdt = child
            break
        # 2) Fallback: field instruction contains TOC
        if any('TOC' in (e.text or '') for e in child.iter(qn('w:instrText'))):
            toc_sdt = child
            break
    if toc_sdt is None:
        return False

    sdt_content = toc_sdt.find(qn('w:sdtContent'))
    if sdt_content is None:
        return False

    ns = nsdecls('w')

    def xe(t):
        return (t.replace('&', '&amp;')
                 .replace('<', '&lt;')
                 .replace('>', '&gt;'))

    # Build entries with section numbers
    entries = []
    first_h1_done = False
    h2_count = 0
    h3_count = 0
    for elem in elements:
        t = elem['type']
        if t == 'h1':
            if not first_h1_done:
                first_h1_done = True
            continue
        if t == 'toc_heading':
            continue
        if t == 'h2':
            h2_count += 1
            h3_count = 0
            entries.append((1, str(h2_count), elem['texto'],
                            elem.get('_bookmark', '')))
        elif t == 'h3':
            h3_count += 1
            entries.append((2, "{}.{}".format(h2_count, h3_count),
                            elem['texto'], elem.get('_bookmark', '')))

    # Clear existing content
    for ch in list(sdt_content):
        sdt_content.remove(ch)

    # -- Title paragraph: "Indice" --
    # Style: Normal (so TOC field never indexes it), but with the exact
    # run and spacing properties of the template H1 style.
    rpr_xml, spc_before, spc_after = _h1_style_xml(doc, settings)

    title_xml = (
        '<w:p {ns}>'
        '<w:pPr>'
          '<w:pStyle w:val="Normal"/>'
          '<w:spacing w:before="{before}" w:after="{after}"/>'
        '</w:pPr>'
        '{rpr}'
        '<w:r>'
          '{rpr}'
          '<w:t xml:space="preserve">\u00cdndice</w:t>'
        '</w:r>'
        '</w:p>'
    ).format(ns=ns, before=spc_before, after=spc_after, rpr=rpr_xml)
    sdt_content.append(parse_xml(title_xml))

    if not entries:
        return True

    for i, (level, num, text, bm) in enumerate(entries):
        is_first = (i == 0)
        is_last  = (i == len(entries) - 1)
        sid = _toc_entry_style(doc, level)

        xml = (
            '<w:p {ns}>'
            '<w:pPr>'
              '<w:pStyle w:val="{sid}"/>'
              '<w:tabs>'
                '<w:tab w:val="right" w:leader="dot" w:pos="9072"/>'
              '</w:tabs>'
            '</w:pPr>'
        ).format(ns=ns, sid=sid)

        if is_first:
            xml += (
                '<w:r><w:fldChar w:fldCharType="begin"/></w:r>'
                '<w:r><w:instrText xml:space="preserve">'
                ' TOC \\o &quot;1-3&quot; \\h \\z \\u '
                '</w:instrText></w:r>'
                '<w:r><w:fldChar w:fldCharType="separate"/></w:r>'
            )

        entry_text = xe("{} {}".format(num, text))
        # Just the plain text inside the TOC field - Word's TablesOfContents.Update()
        # will replace this with properly hyperlinked+numbered entries including
        # correct page numbers, because the field instruction has \h (hyperlinks).
        xml += '<w:r><w:t xml:space="preserve">{}</w:t></w:r>'.format(entry_text)
        xml += '<w:r><w:tab/></w:r><w:r><w:t>1</w:t></w:r>'

        if is_last:
            xml += '<w:r><w:fldChar w:fldCharType="end"/></w:r>'

        xml += '</w:p>'
        sdt_content.append(parse_xml(xml))

    return True


# ---------------------------------------------------------------------------
# Page-break / insertion-point detection
# ---------------------------------------------------------------------------

def find_page_breaks(body):
    breaks = []
    for idx, child in enumerate(list(body)):
        found = False
        for br in child.iter(qn('w:br')):
            if br.get(qn('w:type')) == 'page':
                breaks.append(idx); found = True; break
        if found:
            continue
        for pb in child.iter(qn('w:pageBreakBefore')):
            val = pb.get(qn('w:val'))
            if val is None or val not in ('0', 'false'):
                breaks.append(idx); found = True; break
        if found:
            continue
        if child.tag == qn('w:p'):
            pPr = child.find(qn('w:pPr'))
            if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                breaks.append(idx)
    return breaks

def find_insertion_point(doc, debug=False):
    body = doc.element.body
    children = list(body)

    # Find paragraphs that carry a section-break (w:pPr/w:sectPr).
    # sect_paras[-1] ends the TOC/content section and starts the closing section.
    sect_paras = [
        i for i, ch in enumerate(children)
        if ch.tag == qn('w:p')
        and ch.find(qn('w:pPr')) is not None
        and ch.find(qn('w:pPr')).find(qn('w:sectPr')) is not None
    ]

    if debug:
        print("[DEBUG] Section-break paragraphs at: {}".format(sect_paras))
        print("[DEBUG] Total body children: {}".format(len(children)))

    if len(sect_paras) >= 2:
        # Insert BEFORE the section-break paragraph that starts the closing page.
        # If inserted after it, content belongs to closing section and overlays that page.
        closing_break_para_idx = sect_paras[-1]
        insert_before_elem = children[closing_break_para_idx]
        if debug:
            print("[DEBUG] Inserting before body[{}] (closing page preserved at end)".format(
                closing_break_para_idx))
        return insert_before_elem, []

    elif len(sect_paras) == 1:
        toc_end = sect_paras[0]
        to_remove = [ch for ch in children[toc_end + 1:]
                     if ch.tag != qn('w:sectPr')]
        return None, to_remove

    else:
        to_remove = [
            ch for ch in children
            if ch.tag == qn('w:p')
            and not ''.join(el.text or '' for el in ch.iter()).strip()
        ]
        return None, to_remove


    # ---------------------------------------------------------------------------
    # Content insert helpers
    # ---------------------------------------------------------------------------

def _insert_image_at(doc, ruta, alt, settings, add_paragraph_fn, width_inches=None):
    try:
        if not os.path.exists(ruta):
            print("[AVISO] Image not found: {}".format(ruta)); return
        if width_inches is None:
            width_inches = get_scaled_width(ruta)
        para = add_paragraph_fn()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run().add_picture(ruta, width=Inches(width_inches))
        if alt:
            cap = add_paragraph_fn(style=safe_style(doc, settings, S_CAPTION))
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.add_run(alt)
    except Exception as e:
        print("[AVISO] Error inserting image '{}': {}".format(ruta, e))

def _insert_code_block(doc, codigo, lenguaje, settings, add_paragraph_fn):
    sname = safe_style(doc, settings, S_CODE)
    for line in codigo.split("\n"):
        para = add_paragraph_fn(style=sname)
        run = para.add_run(line)
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        try:
            para._element.get_or_add_pPr().append(
                parse_xml('<w:shd {} w:fill="F0F0F0" w:val="clear"/>'.format(
                    nsdecls('w'))))
        except Exception:
            pass

def _insert_table_at(doc, headers, rows, add_table_fn):
    try:
        nc = len(headers)
        table = add_table_fn(rows=1 + len(rows), cols=nc)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        for ci, h in enumerate(headers):
            cell = table.rows[0].cells[ci]
            cell.text = ""
            aplicar_inline(cell.paragraphs[0], h)
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for ri, row in enumerate(rows):
            for ci, ct in enumerate(row):
                if ci < nc:
                    cell = table.rows[ri + 1].cells[ci]
                    cell.text = ""
                    aplicar_inline(cell.paragraphs[0], ct)
    except Exception as e:
        print("[AVISO] Error inserting table: {}".format(e))


# ---------------------------------------------------------------------------
# Header / Footer
# ---------------------------------------------------------------------------

def set_header(doc, title):
    pass  # template header has logo - do not overwrite


def set_footer(doc, company, title):
    pass  # template footer has logo + page number - do not overwrite


def _set_run_text(run_elem, text):
    t = run_elem.find(qn('w:t'))
    if t is None:
        t = OxmlElement('w:t')
        run_elem.append(t)
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' ' or '  ' in text):
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')


def _set_txbx_para_text(paras, idx, text):
    if idx >= len(paras):
        return
    para = paras[idx]
    runs = para.findall(qn('w:r'))
    if not runs:
        return
    _set_run_text(runs[0], text)
    for extra in runs[1:]:
        para.remove(extra)


def _update_cover(doc, doc_title, settings):
    ns_wps = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    gen     = settings.get('General', {})
    author  = gen.get('Author', '')
    role    = gen.get('AuthorRole', '')
    extra   = gen.get('ExtraInfo', '')
    web_url = gen.get('WebUrl', '')

    body = doc.element.body
    ch0  = list(body)[0]
    wsps = list(ch0.iter('{%s}wsp' % ns_wps))

    def get_content(wsp):
        return wsp.find('.//' + qn('w:txbxContent'))

    # --- wsp[0]: author / extra info ---
    if len(wsps) >= 1:
        c = get_content(wsps[0])
        if c is not None:
            paras = c.findall(qn('w:p'))
            author_line = (author + ' - ' + role) if (author and role) else (author or role)
            _set_txbx_para_text(paras, 0, author_line)
            _set_txbx_para_text(paras, 1, extra)

    # --- wsp[1] title box (wps branch) ---
    if len(wsps) >= 2:
        c = get_content(wsps[1])
        if c is not None:
            paras = c.findall(qn('w:p'))
            if paras:
                para0 = paras[0]
                runs = para0.findall(qn('w:r'))
                if runs:
                    src_idx = None
                    for i, r in enumerate(runs):
                        t = r.find(qn('w:t'))
                        if t is not None and t.text and 'forfuture sans 49 izquierda' in t.text.lower():
                            src_idx = i
                            break
                    if src_idx is None:
                        for i, r in enumerate(runs):
                            rp = r.find(qn('w:rPr'))
                            if rp is not None and rp.find(qn('w:color')) is not None:
                                src_idx = i
                                break
                    if src_idx is None:
                        src_idx = len(runs) - 1
                    _set_run_text(runs[src_idx], doc_title)
                    for i, r in enumerate(list(runs)):
                        if i != src_idx:
                            para0.remove(r)

    # --- Fallback branches (VML/AlternateContent): also replace placeholders ---
    # Some Word renderers show fallback text. Ensure both placeholders are handled:
    #   - clear "Titulo principal"
    #   - replace "ForFuture Sans 49 izquierda" with doc_title
    first_vivid_done = False
    for t in ch0.iter(qn('w:t')):
        val = (t.text or '')
        low = val.lower()
        if ('título principal' in low) or ('titulo principal' in low) or ('t?tulo principal' in low):
            t.text = ''
        elif 'forfuture sans 49 izquierda' in low:
            if not first_vivid_done:
                t.text = doc_title
                first_vivid_done = True
            else:
                t.text = ''

    _update_cover_header(doc, web_url)
    _update_content_footer(doc, doc_title)


def _update_cover_header(doc, web_url):
    # Section 0 header contains "web" text (in a wsp text box) -> replace with WebUrl
    if not web_url:
        return
    ns_wps = 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape'
    try:
        hdr = doc.sections[0].header
        for t in hdr._element.iter(qn('w:t')):
            if t.text and t.text.strip().lower() == 'web':
                t.text = web_url
                return
    except Exception as e:
        print('[AVISO] Could not update cover header: {}'.format(e))


def _update_content_footer(doc, doc_title):
    # Section 1 footer has a hardcoded title text in the template.
    # Replace it with the real document title, preserving original formatting.
    if not doc_title:
        return
    try:
        ftr = doc.sections[1].footer if len(doc.sections) > 1 else None
        if ftr is None:
            return
        for t in ftr._element.iter(qn('w:t')):
            txt = (t.text or '').strip()
            if txt and not txt.isdigit():
                t.text = doc_title
                if doc_title and (' ' in doc_title):
                    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
                return
    except Exception as e:
        print('[AVISO] Could not update content footer: {}'.format(e))

def _update_fields_via_word(docx_path):
    word = None
    try:
        import win32com.client
        import pythoncom
        pythoncom.CoInitialize()
        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        word.DisplayAlerts = 0
        abs_path = os.path.abspath(docx_path)
        doc = word.Documents.Open(
            abs_path,
            ConfirmConversions=False,
            ReadOnly=False,
            AddToRecentFiles=False,
            Visible=False
        )
        doc.Content.Fields.Update()
        if doc.TablesOfContents.Count > 0:
            doc.TablesOfContents(1).Update()
        for section in doc.Sections:
            for header in section.Headers:
                header.Range.Fields.Update()
            for footer in section.Footers:
                footer.Range.Fields.Update()
        doc.Save()
        doc.Close(False)
        word.Quit()
        pythoncom.CoUninitialize()
        print("[OK] Fields updated via Word automation.")
        return True
    except ImportError:
        print("[AVISO] pywin32 not installed. Run: pip install pywin32")
        return False
    except Exception as e:
        print("[AVISO] Could not update fields via Word: {}".format(e))
        try:
            if word:
                word.Quit()
        except Exception:
            pass
        return False


# ---------------------------------------------------------------------------
# Main document generator
# ---------------------------------------------------------------------------

def generate_document(elements, template_path, output_path,
settings, debug=False):
    doc = Document(template_path)
    if debug:
        print("\n[DEBUG] Template styles:")
        for s in doc.styles:
            print("  {} ({})".format(s.name, s.style_id))
        print("\n[DEBUG] Style map: {}".format(settings.get('Styles', {})))

    _assign_bookmarks(elements)

    insert_before, elements_to_remove = find_insertion_point(doc, debug)
    body = doc.element.body
    for elem in elements_to_remove:
        body.remove(elem)

    doc_title = next((e["texto"] for e in elements if e["type"] == "h1"), "")
    mermaid_temps = []
    mermaid_counter = 0

    def _insert_element(element):
        if insert_before is not None:
            insert_before.addprevious(element)
        else:
            body.append(element)

    def _add_paragraph(style=None, text=""):
        from docx.text.paragraph import Paragraph
        p = parse_xml('<w:p {}/>'.format(nsdecls('w')))
        _insert_element(p)
        para = Paragraph(p, doc._body)
        if style:
            try:
                para.style = style
            except Exception:
                pass
        if text:
            para.add_run(text)
        return para

    def _add_table(rows, cols):
        from docx.table import Table
        tbl = doc.add_table(rows=rows, cols=cols)._tbl
        body.remove(tbl)
        _insert_element(tbl)
        return Table(tbl, doc)

    first_h1_skipped = False

    try:
        for elem in elements:
            etype = elem["type"]

            if etype == "h1":
                if not first_h1_skipped:
                    first_h1_skipped = True
                    continue
                para = _add_paragraph(style=safe_style(doc, settings, S_H1))
                aplicar_inline(para, elem["texto"])

            elif etype == "h2":
                para = _add_paragraph(style=safe_style(doc, settings, S_H1))
                aplicar_inline(para, elem["texto"])
                if '_bookmark' in elem:
                    _add_bookmark_to_para(para._element,
                                          elem['_bookmark'], elem['_bm_id'])

            elif etype == "h3":
                para = _add_paragraph(style=safe_style(doc, settings, S_H2))
                aplicar_inline(para, elem["texto"])
                if '_bookmark' in elem:
                    _add_bookmark_to_para(para._element,
                                          elem['_bookmark'], elem['_bm_id'])

            elif etype in ("toc_heading", "toc"):
                pass  # populated below by _update_sdt_toc

            elif etype == "text":
                para = _add_paragraph(style=safe_style(doc, settings, S_BODY))
                aplicar_inline(para, elem["texto"])

            elif etype == "bullet":
                para = _add_paragraph(style=safe_style(doc, settings, S_BULLET))
                aplicar_inline(para, elem["texto"])
                if elem.get("nivel", 0) > 0:
                    try:
                        para._element.get_or_add_pPr().append(
                            parse_xml('<w:ind {} w:left="720"/>'.format(
                                nsdecls('w'))))
                    except Exception:
                        pass

            elif etype == "numbered":
                para = _add_paragraph(style="List Number")
                aplicar_inline(para, elem["texto"])

            elif etype == "code_block":
                _insert_code_block(doc, elem["codigo"],
                                   elem.get("lenguaje", ""),
                                   settings, _add_paragraph)

            elif etype == "image":
                _insert_image_at(doc, elem["ruta"], elem.get("alt", ""),
                                 settings, _add_paragraph)

            elif etype == "mermaid":
                png_path = render_mermaid(elem["codigo"])
                if png_path:
                    mermaid_temps.append(png_path)
                    mermaid_counter += 1
                    _insert_image_at(
                        doc,
                        png_path,
                        "Diagrama Mermaid {}".format(mermaid_counter),
                        settings,
                        _add_paragraph,
                    )
                else:
                    _insert_code_block(doc, elem["codigo"], "mermaid",
                                       settings, _add_paragraph)

            elif etype == "table":
                _insert_table_at(doc, elem["headers"], elem["rows"],
                                 _add_table)

            elif etype == "separator":
                _add_paragraph(style="Normal")

        # Populate SDT TOC - pass settings so title style matches H1
        result = _update_sdt_toc(doc, elements, settings)
        if debug:
            print("[DEBUG] SDT TOC updated: {}".format(result))

        # Update cover page text boxes
        _update_cover(doc, doc_title, settings)

        doc.save(output_path)
        print("[OK] Document generated: {}".format(output_path))

        if sys.platform == "win32":
            _update_fields_via_word(output_path)

    finally:
        for tmp in mermaid_temps:
            try:
                if os.path.exists(tmp):
                    os.unlink(tmp)
            except OSError:
                pass


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main():
    parser = argparse.ArgumentParser(
        description="Convert Markdown to Word with corporate template.")
    parser.add_argument("entrada",   help="Input Markdown file (.md)")
    parser.add_argument("plantilla", help="Word template (.docx)")
    parser.add_argument("salida",    help="Output Word file (.docx)")
    parser.add_argument("--debug",   action="store_true", help="Debug mode")
    args = parser.parse_args()
    if not os.path.exists(args.entrada):
        print("[ERROR] Not found: {}".format(args.entrada)); sys.exit(1)
    if not os.path.exists(args.plantilla):
        print("[ERROR] Not found: {}".format(args.plantilla)); sys.exit(1)
    settings = load_settings()
    if args.debug:
        print("[DEBUG] Input:    {}".format(args.entrada))
        print("[DEBUG] Template: {}".format(args.plantilla))
        print("[DEBUG] Output:   {}".format(args.salida))
    elements = parse_markdown(args.entrada)
    if args.debug:
        print("\n[DEBUG] Parsed elements: {}".format(len(elements)))
        for e in elements:
            print("  {}: {}".format(e['type'], str(e)[:80]))
    generate_document(elements, args.plantilla, args.salida,
                      settings, args.debug)


if __name__ == "__main__":
    main()
